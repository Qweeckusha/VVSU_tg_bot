'''
                             _    ___    __________  __     __          __
                            | |  / / |  / / ____/ / / /    / /_  ____  / /_
                            | | / /| | / / / __/ / / /    / __ \/ __ \/ __/
                            | |/ / | |/ / /_/ / /_/ /    / /_/ / /_/ / /_
                            |___/  |___/\____/\____/____/_.___/\____/\__/
                                                  /_____/

if Qweeck:
    print('Холохон Анатолий Константинович')
    Qweeck.status = 'dev'

'''
from datetime import datetime
import pymorphy2
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters import Command, Text
from io import BytesIO
from aiogram import Bot, Dispatcher, types
from aiogram.dispatcher.filters.state import StatesGroup, State
from aiogram.types import InputFile, ReplyKeyboardMarkup, KeyboardButton, InlineKeyboardButton, InlineKeyboardMarkup
from aiogram import executor
from docx import Document
from docx.shared import Pt


from config import API_TOKEN

storage = MemoryStorage()
bot = Bot(token=API_TOKEN)
dp = Dispatcher(bot, storage=storage)



# -------------------------------------------------------------------
# Классы для каждого документа, относительно вводимого в него текста.
# -------------------------------------------------------------------
class DocFormStates(StatesGroup):
    director_full_name = State()
    group = State()
    full_name = State()
    contact = State()
    date_from = State()
    date_to = State()
    org_name = State()
    org_phone = State()
    full_name_org = State()
    org_job = State()
    full_name_kafedra = State()
    reasons = State()
    date_to2 = State()
    legal_adress = State()
    mail = State()
    inn = State()
    ogrn = State()
    okpo = State()
    bik = State()
    direction = State()
    type_practice = State()
    date_from_pr = State()
    date_to_pr = State()



@dp.message_handler(Command("start"))
async def process_start_command(message: types.Message):
    markup = ReplyKeyboardMarkup(resize_keyboard=True, keyboard=[
        [KeyboardButton('Создать заявление и договор на практику')]
    ])
    await message.answer("Привет!\nЯ бот для создания заявления и договора о практической подготовки.", reply_markup=markup)
    await message.answer('Чтобы начать заполнение заявления и договора нажми кнопку снизу в меню.',
                         reply_markup=InlineKeyboardMarkup(
                             inline_keyboard=[
                                [InlineKeyboardButton('Где меню?', callback_data='wheremenu')]
    ]))
    await message.delete()

#new code
@dp.callback_query_handler(lambda c: c.data == 'wheremenu')
async def where_menu(call: types.CallbackQuery):
        await call.bot.send_photo(call.message.chat.id, photo=InputFile('templates/menu.jpg'), caption='Кнопка меню находится рядом со строкой ввода сообщения (справа).')
# ------------------------------------------------------
# ↓↓↓ Алгоритм сбора данных для заполнения документа "Заявление на практику" ↓↓↓
# ------------------------------------------------------

@dp.message_handler(Text(equals='Создать заявление и договор на практику'))
async def process_personal_data(message: types.Message):

    await message.answer('❗Согласие на обработку персональных данных❗\n\nВы согласны на '
                         'предоставление своих персональных данных?',parse_mode='HTML', reply_markup=InlineKeyboardMarkup(
        inline_keyboard=[
        [InlineKeyboardButton('✅ Да ✅', callback_data='PDyes'), InlineKeyboardButton('❌ Нет ❌', callback_data='PDno')],
        [InlineKeyboardButton('📃 Согласие на обработку ПДн 📃', callback_data='PDcheck')]

    ]), disable_web_page_preview = True)

@dp.callback_query_handler(lambda c: c.data.startswith('PD'))
async def process_create_document(call: types.CallbackQuery):
    if call.data == 'PDyes':
        markup = ReplyKeyboardMarkup(resize_keyboard=True, keyboard=[
            [KeyboardButton('Отмена')]
        ])
        # Запускаем состояние
        await DocFormStates.group.set()
        await call.message.answer('❗Обращаю внимание❗\nПеред отправкой проверяйте корректность введённых данных.\n\n'
                             'Для прекращения заполнения документов нажмите кнопку "Отмена" снизу. Можете '
                             'использовать это в качестве меры предотвращения опечатки. \n\nВ любом случае готовый '
                             'документ тоже можно отредактировать.', reply_markup=markup)
        await call.message.answer("<b>Сейчас вы заполняете заявление на закрепление места в практике!</b>", parse_mode='HTML')
        await call.message.answer("Введите номер вашей группы, например: <b>БИН-23-1</b>", parse_mode='HTML')

    elif call.data == 'PDcheck':
        await call.bot.send_document(call.message.chat.id, InputFile('templates/Согласие_на_обработку_ПДн.docx'),
                                     caption='Этот файл будет внутри договора на практическую подготовку.')
    else:
        await call.message.answer('Вы не согласны на обработку персональных данных, заполнение невозможно.')
@dp.message_handler(lambda message: 'отмена' in message.text.lower(), state="*")
async def cancel_handler(message: types.Message, state: FSMContext):
    current_state = await state.get_state()
    if current_state is None:
        return
    await state.finish()
    markup = ReplyKeyboardMarkup(resize_keyboard=True, keyboard=[
        [KeyboardButton('Создать заявление и договор на практику')]
    ])
    await message.answer('Заполнение формы отменено. Вы можете начать заново.', reply_markup=markup)

@dp.message_handler(state=DocFormStates.group)
async def process_group(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['group'] = message.text


    await DocFormStates.full_name.set()
    await message.answer("Введите ваше ФИО, например: <b>Иванов Иван Иванович</b>", parse_mode='HTML')

@dp.message_handler(state=DocFormStates.full_name)
async def process_full_name(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['full_name'] = message.text


    await DocFormStates.director_full_name.set()
    await message.answer("Введите ФИО директора института/заведующего кафедрой, например: <b>Иванов Иван Иванович</b>", parse_mode='HTML')

@dp.message_handler(state=DocFormStates.director_full_name)
async def process_full_name_director(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['full_name_director'] = message.text


    await DocFormStates.contact.set()
    await message.answer("Введите свои контакты номер телефона <b>ИЛИ</b> e-mail, например <b>+79123458796 | 89123458796</b> или <b>email@gmail.com</b>", parse_mode='HTML')

@dp.message_handler(state=DocFormStates.contact)
async def process_contact(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['contact'] = message.text
    await DocFormStates.date_from.set()
    await message.answer("Введите дату начала практики,"
                         " например: <b>24 апреля 2024</b>", parse_mode='HTML', reply_markup=types.InlineKeyboardMarkup(inline_keyboard=[
        [types.InlineKeyboardButton('Если непонятно', callback_data='desc')]
    ]))

@dp.callback_query_handler(lambda c: c.data == 'desc', state=DocFormStates.date_from)
async def call_desc(call: types.CallbackQuery):
    await call.message.answer('Вам нужно ввести дату начала вашей первой практики, например:\n <b>Ваша первая '
                              'практика учебная, следовательно вы должны использовать здесь '
                              'дату начала учебной практики</b>\n\n❗Идея в том, что вы должны указать начало первой '
                              'практики и конец самой последней (ЭТО СЛЕДУЮЩИЙ ВОПРОС), '
                              'из этого складывается весь период прохождения практики в течение учебного процесса', parse_mode='HTML')

@dp.message_handler(state=DocFormStates.date_from)
async def process_date_from(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['date_from'] = message.text

    await DocFormStates.date_to.set()
    await message.answer("Введите дату окончания последней практики, например: <b>15 июля 2025</b>", parse_mode='HTML')

@dp.message_handler(state=DocFormStates.date_to)
async def process_date_to(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['date_to'] = message.text

    await DocFormStates.org_name.set()
    await message.answer('Введите наименование организации, например: <b>ООО "Юником"</b>', parse_mode='HTML')

@dp.message_handler(state=DocFormStates.org_name)
async def process_org_name(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['org_name'] = message.text

    await DocFormStates.org_phone.set()
    await message.answer("Введите контактный номер организации, например: <b>+78005553535 или 88005553535</b>", parse_mode='HTML')
# Эти данные пригодятся и в договоре


@dp.message_handler(state=DocFormStates.org_phone)
async def process_org_phone(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['org_phone'] = message.text

    await DocFormStates.full_name_org.set()
    await message.answer("Введите ФИО руководителя от организации, например: <b>Иванов Иван Иванович</b>", parse_mode='HTML')

@dp.message_handler(state=DocFormStates.full_name_org)
async def process_full_name_org(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['full_name_org'] = message.text

    await DocFormStates.org_job.set()
    await message.answer("Введите его должность, например: <b>управляющий директор</b>", parse_mode='HTML')

@dp.message_handler(state=DocFormStates.org_job)
async def process_org_job(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['org_job'] = message.text.capitalize()


    await fill_document(
        template_path='templates/Заявление_на_практику.docx',
        director_full_name=data['full_name_director'],
        group=data['group'],
        full_name=data['full_name'],
        contact=data['contact'],
        date_from=data['date_from'],
        date_to=data['date_to'],
        org_name=data['org_name'],
        org_phone=data['org_phone'],
        full_name_org=data['full_name_org'],
        org_job=data['org_job'].capitalize(),
        f_initial=format_fio(data['full_name']),
        message=message
    )

    await message.answer("<b>Заявление заполнено!</b>",
                         parse_mode='HTML')
    await DocFormStates.reasons.set()
    await message.answer("Напишите, на основании чего действует организация, например: <b>устава или доверенности №</b>",
                         parse_mode='HTML')

# -------------------------------------------------------------------------------
# ↓↓↓ Алгоритм для сбора данных документа "Договор о практической подготовке" ↓↓↓
# -------------------------------------------------------------------------------


@dp.message_handler(state=DocFormStates.reasons)
async def process_reasons(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['reasons'] = message.text.lower()

        await DocFormStates.date_to2.set()
        await message.answer("Напишите, до какого числа будет действителен договор, например: <b>24.02.2024</b>",
                             parse_mode='HTML')

@dp.message_handler(state=DocFormStates.date_to2)
async def process_dateto2(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['date_to2'] = message.text

        await DocFormStates.legal_adress.set()
        await message.answer("Введите юридический адрес компании, например: <b>690008, г. Владивосток, ул. Державина, д.14, кв. 57</b>",
                             parse_mode='HTML')

@dp.message_handler(state=DocFormStates.legal_adress)
async def process_legal_adress(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['legal_adress'] = message.text

        await DocFormStates.mail.set()
        await message.answer("Введите почтовый (фактический) адрес организации, например: <b>690003, г. Владивосток, ул. Луговая, д.17</b>",
                             parse_mode='HTML')

@dp.message_handler(state=DocFormStates.mail)
async def process_mail(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['mail'] = message.text

        await DocFormStates.inn.set()
        await message.answer("Введите ИНН или КПП организации, например: <b>121212121212</b>",
                                parse_mode='HTML')

@dp.message_handler(state=DocFormStates.inn)
async def process_inn(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['inn'] = message.text

        await DocFormStates.ogrn.set()
        await message.answer("Введите ОГРН организации, например: <b>151515151515151</b>",
                             parse_mode='HTML')

@dp.message_handler(state=DocFormStates.ogrn)
async def process_ogrn(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['ogrn'] = message.text

        await DocFormStates.okpo.set()
        await message.answer('Введите ОКПО организации (если он отсутствует, то отправьте "Нет"), например: <b>88888888</b>',
                             parse_mode='HTML')

@dp.message_handler(state=DocFormStates.okpo)
async def process_okpo(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        if message.text == '-' or message.text == 'Нет' or message.text == 'нет':
            data['okpo'] = ''
        else:
            data['okpo'] = f'ОКПО: {message.text}'

    await DocFormStates.bik.set()
    await message.answer("Введите БИК организации (9 цифр), например: <b>999999999</b>",
                             parse_mode='HTML')

@dp.message_handler(state=DocFormStates.bik)
async def process_bik(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['bik'] = message.text

    await DocFormStates.direction.set()
    await message.answer("Укажите ваше направление подготовки, например: <b>Программная инженерия</b>", parse_mode='HTML')

@dp.message_handler(state=DocFormStates.direction)
async def process_direction(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['direction'] = message.text

    await DocFormStates.type_practice.set()
    await message.answer("Выберите тип практики или напишите свой вариант. <b>Используйте меню снизу.</b>", parse_mode='HTML',
                         reply_markup=ReplyKeyboardMarkup(resize_keyboard=True, keyboard=[
        [KeyboardButton('Учебная')], [KeyboardButton('Производственная')], [KeyboardButton('Преддипломная')]
    ]))

@dp.message_handler(state=DocFormStates.type_practice)
async def process_type_practice(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        if 'type_practice' not in data:
            data['type_practice'] = '' # Создаём переменную, чтобы можно было переприсвоить
        if data['type_practice'] == '':
            data['type_practice'] = data['type_practice'] + message.text
        else:
            data['type_practice'] = data['type_practice'] + '\n\n' + message.text

        await DocFormStates.date_from_pr.set()
        await message.answer("Введите дату её начала, например: <b>24.03.2024</b>", parse_mode='HTML')

@dp.message_handler(state=DocFormStates.date_from_pr)
async def process_date_from_pr(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        data['date_from_pr'] = message.text

    await DocFormStates.date_to_pr.set()
    await message.answer("И укажите дату её окончания, например: <b>21.04.2024</b>", parse_mode='HTML')

@dp.message_handler(state=DocFormStates.date_to_pr)
async def process_date_to_pr(message: types.Message, state: FSMContext):
    async with state.proxy() as data:
        if message.text not in ['Дополнить', 'Нет']:
            if 'date_complete_pr' not in data:
                data['date_complete_pr'] = ''
            if data['date_complete_pr'] == '':
                data['date_complete_pr'] = data['date_complete_pr'] + f'{data["date_from_pr"]}-{message.text}'
            else:
                data['date_complete_pr'] = data['date_complete_pr'] + '\n\n' + f'{data["date_from_pr"]}-{message.text}'

            markup = ReplyKeyboardMarkup(resize_keyboard=True, keyboard=[
                [KeyboardButton('Дополнить')], [KeyboardButton('Нет')]
            ])
            await message.answer('Если у вас несколько периодов практики, то нажмите '
                                 '"Дополнить", иначе нажмите "Нет"', parse_mode='HTML', reply_markup=markup)

        if message.text.lower() == 'дополнить':
            await DocFormStates.type_practice.set()
            await message.answer("Выберите тип практики или напишите свой вариант. <b>Используйте меню снизу.</b>",
                                 parse_mode='HTML',
                                 reply_markup=ReplyKeyboardMarkup(resize_keyboard=True, keyboard=[
                [KeyboardButton('Учебная')], [KeyboardButton('Производственная')], [KeyboardButton('Преддипломная')]
            ]))
        elif message.text.lower() == 'нет':
            markup = ReplyKeyboardMarkup(resize_keyboard=True, keyboard=[
                [KeyboardButton('Создать заявление и договор на практику')]
            ])
            await fill_document2(
                template_path='templates/Договор_ВВГУ_о_практической_подготовке.docx',
                organization=data['org_name'],
                position=f"{get_genitive(data['org_job'])} {get_genitive_cap(data['full_name_org'])}",
                reasons=first_word_lower(data['reasons']),
                date_to2=data['date_to2'],
                legal_adress=data['legal_adress'],
                mail=data['mail'],
                inn=data['inn'],
                ogrn=data['ogrn'],
                okpo=data['okpo'],
                bik=data['bik'],
                contact=data['org_phone'],
                profession=data['org_job'].capitalize(),
                ceo=format_fio(data['full_name_org']),
                full_name=data['full_name'],
                phone=data['contact'],
                direction=data['direction'],
                group=data['group'],
                date_complete=data['date_complete_pr'],
                type_practice=data['type_practice'],
                message=message
            )
            await message.answer('Вы завершили заполнение последнего документа!🥳🥳🥳', reply_markup=markup)
            await message.answer('❗❗❗ВНИМАНИЕ❗❗❗\n<b>Проверьте заявление  и договор на правильность перед печатью.</b>\n\n'
                                 '1. Не забудьте <b>распечатать и вручную</b> заполнить согласие на обработку персональных '
                                 'данных (оно находится файле "Согласие на обработку ПДн.docx").\n\n'
                                 '2. Распечатайте в 2-х экземплярах договор, отнесите их в "Старт-Карьеру" (аудитория 1442).\n\n'
                                 '3. Распечатайте заявление и отнесите в свою кафедру.', parse_mode='HTML')
            await state.finish()

def format_fio(full_name):
    fio_parts = full_name.split()
    initials = ' '.join([name[0] + '.' for name in fio_parts[1:]])
    return f"{fio_parts[0]} {initials}"

def first_word_lower(stroka):
    if len(stroka.split()) > 1:
        words = stroka.split()
        other_words = ' '.join([name for name in words[1:]])
        return f'{words[0].lower()} {other_words}'
    else:
        return stroka.lower()

def get_genitive(words):
    try:
        morph = pymorphy2.MorphAnalyzer()
        return ' '.join([morph.parse(i)[0].inflect({'gent'}).word for i in words.split()])
    except:
        return words


def get_genitive_cap(words):
    try:
        morph = pymorphy2.MorphAnalyzer()
        return ' '.join([morph.parse(i)[0].inflect({'gent'}).word.capitalize() for i in words.split()])
    except:
        return words

# ---------------------------------------------------------
# ↓↓↓ Алгоритм для внесения в документ собранных данных ↓↓↓
# ---------------------------------------------------------

async def fill_document(
        template_path,
        director_full_name,
        group,
        full_name,
        contact,
        date_from,
        date_to,
        org_name,
        org_phone,
        full_name_org,
        org_job,
        f_initial,
        message
        ):
    # Открываем существующий документ
    doc = Document(template_path)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    # Заменяем метки в документе на соответствующие данные
    replace_text(doc, "{{FULL_NAME_DIRECTOR}}", director_full_name)
    replace_text(doc, "{{GROUP}}", group, underline=True)
    replace_text(doc, "{{FULL_NAME}}", full_name, underline=True)
    replace_text(doc, "{{CONTACT}}", contact, underline=True)
    replace_text(doc, "{{DATE_FROM}}", date_from, space=False)
    replace_text(doc, "{{DATE_TO}}", date_to, space=False)
    replace_text(doc, "{{ORG_NAME}}", org_name)
    replace_text(doc, "{{ORG_PHONE}}", org_phone)
    replace_text(doc, "{{FULL_NAME_ORG}}", full_name_org, underline=True)
    replace_text(doc, "{{ORG_JOB}}", org_job, underline=True)
    replace_text(doc, "{{TODAY_DATE}}", datetime.now().strftime('«%d» %m %Y'), space=False)
    replace_text(doc, "{{INITIALS_STUD}}", f_initial, space=False)




    # Сохраняем результат
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    await bot.send_document(message.chat.id, InputFile(doc_bytes, filename='Заявление.docx'))

def replace_text(doc, placeholder, new_text, underline=False, space=True):
    placeholders = ['{{FULL_NAME_DIRECTOR}}', '{{ORG_PHONE}}', '{{ORG_NAME}}']
    for paragraph in doc.paragraphs:
        if placeholder in paragraph.text:
            # Заменяем метку на новый текст

            if placeholder in placeholders:
                paragraph.text = paragraph.text.replace(placeholder, '')
                paragraph.add_run('    ' + new_text + '    ').underline = True
            elif placeholder == '{{INITIALS_STUD}}':
                paragraph.text = paragraph.text.replace(placeholder, '')
                paragraph.add_run(new_text).underline = True
                paragraph.add_run(' (инициалы, фамилия)').font.size = Pt(10)
            else:
                if space:
                    paragraph.text = paragraph.text.replace(placeholder, '    ' + new_text + '    ')
                    if underline:
                        for run in paragraph.runs:
                            run.underline = True
                else:
                    paragraph.text = paragraph.text.replace(placeholder, new_text)
                    if underline:
                        for run in paragraph.runs:
                            run.underline = True

# ---------------------------------------------------------
# ↓↓↓ Алгоритм для внесения в документ 2 собранных данных ↓↓↓
# ---------------------------------------------------------
async def fill_document2(
        template_path,
        organization,
        position,
        reasons,
        date_to2,
        legal_adress,
        mail,
        inn,
        ogrn,
        okpo,
        bik,
        contact,
        profession,
        ceo,
        full_name,
        phone,
        direction,
        group,
        date_complete,
        type_practice,
        message
        ):
    # Открываем существующий документ
    doc = Document(template_path)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Заменяем метки в документе на соответствующие данные
    replace_text_special(doc, organization, position, reasons, date_to2)
    replace_text2(doc, '{{ORG_NAME}}', organization)
    replace_text2(doc, '{{LEGAL_ADRESS}}', legal_adress)
    replace_text2(doc, '{{MAIL}}', mail)
    replace_text2(doc, '{{INN}}', inn)
    replace_text2(doc, '{{OGRN}}', ogrn)
    replace_text2(doc, '{{OKPO}}', okpo)
    replace_text2(doc, '{{BIK}}', bik)
    replace_text2(doc, '{{CONTACT}}', contact)
    replace_text2(doc, '{{PROFESSION}}', profession)
    replace_text2(doc, '{{CEO}}', ceo)
    replace_text2(doc, '{{FULL_NAME}}', full_name)
    replace_text2(doc, '{{PHONE}}', phone)
    replace_text2(doc, '{{DIRECTION}}', direction)
    replace_text2(doc, '{{GROUP}}', group)
    replace_text2(doc, '{{DATE_COMPLETE}}', date_complete)
    replace_text2(doc, '{{MAIL}}', mail)
    replace_text2(doc, '{{TYPE_PRACTICE}}', type_practice)

    # Сохраняем результат
    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    await bot.send_document(message.chat.id, InputFile(doc_bytes, filename=f'Договор.docx'))
    await bot.send_document(message.chat.id, InputFile('templates/Согласие_на_обработку_ПДн.docx', filename=f'Согласие на обработку ПДн.docx'))



def replace_text_special(doc, new_text, new_text2, new_text3, new_text4):
    for paragraph in doc.paragraphs:
        if '{{ABOUT_ORG}}' in paragraph.text:
            # Заменяем метку на новый текст

            paragraph.text = paragraph.text.replace('{{ABOUT_ORG}}', '')
            paragraph.add_run(new_text).underline = True
            paragraph.add_run(', именуемое в дальнейшем «Профильная организация», в лице ')

            paragraph.add_run(new_text2).underline = True
            paragraph.add_run(', действующего на основании ')

            paragraph.add_run(new_text3).underline = True
            paragraph.add_run(', с другой стороны, именуемые по '
                                     'отдельности «Сторона», а вместе – «Стороны», заключили настоящий Договор о '
                                     'нижеследующем.')

        if '{{DATE_TO}}' in paragraph.text:
            paragraph.text = paragraph.text.replace('{{DATE_TO}}', '')
            paragraph.add_run(new_text4).underline = True
            paragraph.add_run(' и автоматически продлевается на один год, если ни одна из сторон за 30 (тридцать) '
                              'рабочих дней не заявит о своем намерении его расторгнуть.')


def replace_text2(doc, placeholder, new_text):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        # Заменяем метку на новый текст
                        paragraph.text = paragraph.text.replace(placeholder, new_text)
                        # Устанавливаем стиль текста после замены
                        for run in paragraph.runs:
                            run.font.size = Pt(12)  # Размер шрифта 12
                            run.font.bold = False   # Убираем жирное выделение



try:
    executor.start_polling(dp, skip_updates=True)
except:
    # async def error(message: types.Message):
    #     markup = ReplyKeyboardMarkup(resize_keyboard=True, keyboard=[
    #         [KeyboardButton('Создать заявление и договор на практику')]
    #     ])
    #     await message.answer('Ошибка, попробуйте снова. /start', reply_markup=markup)
    print('Ошибочка.')